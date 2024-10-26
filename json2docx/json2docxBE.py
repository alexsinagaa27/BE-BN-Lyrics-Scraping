import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter import Tk
from tkinter.filedialog import askopenfilename
from docx.shared import RGBColor


def json_to_word(json_data, output_filename="Buku Ende HKBP (Uppercase).docx"):
    # Membuat dokumen Word baru
    doc = Document()
    
    # Menambahkan judul di header pada halaman pertama
    header = doc.sections[0].header
    paragraph = header.add_paragraph()
    run = paragraph.add_run("Buku Ende HKBP")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(128, 0, 0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Iterasi setiap lagu dalam JSON
    for lagu in json_data:
        # Menambahkan judul lagu dengan format "BE NO. {nomorBE} \"{judul}\"" dalam uppercase
        doc.add_heading(f"BE NO. {lagu['nomorBE']} \"{lagu['judul'].upper()}\"", level=1)  # Menambahkan heading
        
        # Menambahkan lirik per bait
        for bait in lagu.get("baits", []):
            # Nomor bait dengan format hanya kurung tutup
            doc.add_paragraph(f"{bait['nomor_bait']})", style="BodyText")  # Format nomor bait
            
            # Menambahkan setiap baris lirik dalam uppercase
            for baris in bait["lirik"]:
                doc.add_paragraph(baris.upper(), style="BodyText")  # Baris lirik dalam uppercase
            
            # doc.add_paragraph("")  # Tambahkan spasi antar bait
        
        # Tambahkan satu baris kosong setelah setiap lagu
        # doc.add_paragraph("")  # Jarak antar lagu
        
    # Simpan dokumen ke file Word
    doc.save(output_filename)
    print(f"Dokumen berhasil disimpan sebagai {output_filename}")

# Menggunakan Tkinter untuk memilih file JSON
def choose_json_file():
    Tk().withdraw()  # Menyembunyikan jendela utama Tkinter
    filename = askopenfilename(filetypes=[("JSON files", "*.json")])
    return filename

# Meminta pengguna memilih file JSON
json_file = choose_json_file()

# Membaca data JSON dari file yang dipilih
if json_file:  # Pastikan file tidak kosong
    with open(json_file, "r", encoding="utf-8") as f:
        json_data = json.load(f)

    # Mengonversi JSON ke Word
    json_to_word(json_data)
else:
    print("Tidak ada file yang dipilih.")
