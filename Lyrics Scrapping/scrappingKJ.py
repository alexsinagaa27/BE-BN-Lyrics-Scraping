import requests # untuk send request dan recieve respons dari website tertentu
from bs4 import BeautifulSoup # untuk parsing HTML agar dapat mengekstrak elemen tertentu
import json # format umum untuk menyimpan data dalam bentuk semi-structured

from docx import Document # Membuat objek dokumen Word baru
from docx.shared import Pt # Menentukan ukuran font dalam point
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT # Mengatur perataan paragraf
from docx.shared import RGBColor # Mengatur warna font menggunakan nilai RGB

def scrape_and_save_to_json(output_filename="data_laguKJ.json"):
    # List untuk menyimpan hasil sementara
    data_lagu = []

    # Loop dari 1 sampai 478
    for nomorKJ in range(1, 479):
        # Membuat URL dinamis dengan nomorKJ
        url = f"https://alkitab.app/KJ/{nomorKJ}"
        
        # Mengirim permintaan untuk mendapatkan halaman
        response = requests.get(url)
        
        # Memeriksa apakah permintaan berhasil
        if response.status_code == 200:
            # Parsing konten HTML
            soup = BeautifulSoup(response.content, "html.parser")
            
            # Mengambil judul
            judul_div = soup.find("div", class_="judul")
            judul = judul_div.get_text().strip() if judul_div else "Judul tidak ditemukan"
            
            # Mengambil nada dasar
            nadaDasar_div = soup.find("div", class_="nadaDasar")
            nadaDasar = nadaDasar_div.get_text().strip() if nadaDasar_div else "Nada Dasar tidak ditemukan"

            # Mengambil lirik per bait
            bait_list = []
            bait_divs = soup.find_all("div", class_="bait")
            for bait in bait_divs:
                # Ambil nomor bait
                bait_no_div = bait.find("div", class_="bait-no")
                bait_no = bait_no_div.get_text().strip() if bait_no_div else "Tidak ada nomor bait"
                
                # Ambil semua baris dalam bait
                baris_lirik = [baris.get_text().strip() for baris in bait.find_all("div", class_="baris")]
                
                # Gabungkan nomor bait dan liriknya
                bait_list.append({
                    "nomor_bait": bait_no,
                    "lirik": baris_lirik
                })
            
            # Menyimpan data dalam dictionary
            lagu = {
                "nomorKJ": nomorKJ,
                "judul": judul,
                "nada_dasar": nadaDasar,
                "baits": bait_list
            }
            
            # Tambahkan data lagu ke dalam list
            data_lagu.append(lagu)

        # Simpan data setiap 25 iterasi
        if nomorKJ % 25 == 0:
            try:
                # Membaca data dari file sebelumnya jika ada
                with open(output_filename, "r", encoding="utf-8") as f:
                    all_data = json.load(f)
            except (FileNotFoundError, json.JSONDecodeError):
                all_data = []  # Jika file belum ada atau kosong

            # Tambahkan data baru ke data yang sudah ada
            all_data.extend(data_lagu)
            
            # Tulis kembali semua data ke file JSON
            with open(output_filename, "w", encoding="utf-8") as f:
                json.dump(all_data, f, ensure_ascii=False, indent=4)
            print(f"Progress tersimpan pada iterasi {nomorKJ}")
            
            # Kosongkan data_lagu untuk iterasi selanjutnya
            data_lagu = []

    # Simpan data yang tersisa (jika ada) di akhir
    if data_lagu:
        try:
            with open(output_filename, "r", encoding="utf-8") as f:
                all_data = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            all_data = []
        
        all_data.extend(data_lagu)
        
        with open(output_filename, "w", encoding="utf-8") as f:
            json.dump(all_data, f, ensure_ascii=False, indent=4)
        print(f"Data tersisa disimpan ke file {output_filename}")

def json_to_word(json_data, output_filename="KIDUNG_JEMAAT.docx"):
    # Membuat dokumen Word baru
    doc = Document()

    # Menambahkan judul di header pada halaman pertama
    header = doc.sections[0].header
    paragraph = header.add_paragraph()
    run = paragraph.add_run("Kidung Jemaat")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(128, 0, 0)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Iterasi setiap lagu dalam JSON
    for lagu in json_data:
        # Menambahkan judul lagu
        doc.add_heading(f"KJ NO. {lagu['nomorKJ']} \"{lagu['judul'].upper()}\"", level=1)

        # Menambahkan nada dasar
        doc.add_paragraph(f"Nada Dasar: {lagu.get('nada_dasar', 'Tidak tersedia')}")

        # Menambahkan lirik per bait
        for bait in lagu.get("baits", []):
            doc.add_paragraph(f"{bait['nomor_bait']})", style="BodyText")
            for baris in bait["lirik"]:
                doc.add_paragraph(baris.upper(), style="BodyText")

        # Tambahkan satu baris kosong setelah setiap lagu
        doc.add_paragraph("")

    # Simpan dokumen Word
    doc.save(output_filename)
    print(f"Dokumen berhasil disimpan sebagai {output_filename}")

# Scraping data dan menyimpannya ke JSON
scrape_and_save_to_json()

# Membaca data JSON
with open("data_laguKJ.json", "r", encoding="utf-8") as f:
    json_data = json.load(f)

# Mengonversi data JSON ke dokumen Word
json_to_word(json_data)
